from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.oxml.shared import OxmlElement as SharedOxmlElement
from docx.shared import Inches, Pt, RGBColor


OUTPUT_DIR = Path(__file__).resolve().parents[1] / "legal"
OUTPUT_DOCX = OUTPUT_DIR / "Crout_Automations_Independent_Contractor_Agreement_Template.docx"

ACCENT = RGBColor(14, 46, 74)
ACCENT_DARK = RGBColor(8, 30, 49)
GOLD = RGBColor(165, 123, 36)
TEXT = RGBColor(24, 32, 38)
MUTED = RGBColor(92, 102, 112)
BORDER = "D7DEE5"


def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge in ("top", "left", "bottom", "right"):
        edge_data = kwargs.get(edge)
        if not edge_data:
            continue
        tag = f"w:{edge}"
        element = tc_borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            tc_borders.append(element)
        for key, value in edge_data.items():
            element.set(qn(f"w:{key}"), str(value))


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def add_page_number(paragraph):
    run = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char_separate = OxmlElement("w:fldChar")
    fld_char_separate.set(qn("w:fldCharType"), "separate")
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char_begin, instr_text, fld_char_separate, fld_char_end])


def add_toc(paragraph):
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = r'TOC \o "1-3" \h \z \u'
    fld_separate = OxmlElement("w:fldChar")
    fld_separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "Right-click and update the table of contents in Word if page numbers do not refresh automatically."
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_begin, instr, fld_separate, text, fld_end])


def force_font(run, name="Aptos", size=11, color=TEXT, bold=False, italic=False):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.bold = bold
    run.italic = italic


def set_para_format(paragraph, before=0, after=6, line=1.12, alignment=WD_ALIGN_PARAGRAPH.LEFT):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line
    paragraph.alignment = alignment


def add_styled_paragraph(doc, text="", style=None, before=0, after=6, line=1.12, alignment=WD_ALIGN_PARAGRAPH.LEFT):
    p = doc.add_paragraph(style=style)
    set_para_format(p, before=before, after=after, line=line, alignment=alignment)
    if text:
        run = p.add_run(text)
        force_font(run)
    return p


def add_run_text(paragraph, text, **kwargs):
    run = paragraph.add_run(text)
    force_font(run, **kwargs)
    return run


def set_doc_language(document, lang="en-ZA"):
    styles = document.styles
    doc_defaults = styles.element.find(qn("w:docDefaults"))
    if doc_defaults is None:
        return
    rpr_default_wrapper = doc_defaults.find(qn("w:rPrDefault"))
    if rpr_default_wrapper is None:
        return
    rpr_default = rpr_default_wrapper.find(qn("w:rPr"))
    if rpr_default is None:
        return
    lang_el = rpr_default.find(qn("w:lang"))
    if lang_el is None:
        lang_el = OxmlElement("w:lang")
        rpr_default.append(lang_el)
    lang_el.set(qn("w:val"), lang)


def configure_styles(document):
    styles = document.styles

    normal = styles["Normal"]
    normal.font.name = "Aptos"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Aptos")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Aptos")
    normal.font.size = Pt(11)
    normal.font.color.rgb = TEXT
    pf = normal.paragraph_format
    pf.space_after = Pt(6)
    pf.line_spacing = 1.12

    for style_name, size, color, before, after in (
        ("Heading 1", 15, ACCENT, 16, 8),
        ("Heading 2", 12.5, ACCENT, 10, 4),
        ("Heading 3", 11.5, ACCENT_DARK, 8, 3),
    ):
        style = styles[style_name]
        style.font.name = "Aptos"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Aptos")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Aptos")
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.08

    if "Clause List" not in styles:
        style = styles.add_style("Clause List", WD_STYLE_TYPE.PARAGRAPH)
        style.base_style = styles["List Number"]
    style = styles["Clause List"]
    style.font.name = "Aptos"
    style._element.rPr.rFonts.set(qn("w:ascii"), "Aptos")
    style._element.rPr.rFonts.set(qn("w:hAnsi"), "Aptos")
    style.font.size = Pt(11)
    style.font.color.rgb = TEXT
    style.paragraph_format.space_after = Pt(5)
    style.paragraph_format.line_spacing = 1.12
    style.paragraph_format.left_indent = Inches(0.25)

    if "Small Caps Label" not in styles:
        style = styles.add_style("Small Caps Label", WD_STYLE_TYPE.PARAGRAPH)
    style = styles["Small Caps Label"]
    style.font.name = "Aptos"
    style._element.rPr.rFonts.set(qn("w:ascii"), "Aptos")
    style._element.rPr.rFonts.set(qn("w:hAnsi"), "Aptos")
    style.font.size = Pt(9.5)
    style.font.color.rgb = MUTED
    style.font.small_caps = True
    style.paragraph_format.space_after = Pt(2)
    style.paragraph_format.line_spacing = 1.0


def set_section_layout(section):
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.49)
    section.footer_distance = Inches(0.49)


def add_footer(section):
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_para_format(p, after=0, line=1.0)
    add_run_text(p, "Crout Holdings (Pty) Ltd | Independent Contractor / Subcontractor Agreement | Page ", size=9.5, color=MUTED)
    add_page_number(p)


def add_header(section):
    header = section.header
    p = header.paragraphs[0]
    set_para_format(p, after=0, line=1.0)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    add_run_text(p, "Crout Automations", size=9.5, color=ACCENT, bold=True)
    add_run_text(p, " | Contractor Template", size=9.5, color=MUTED)
    p_bdr = p._p.get_or_add_pPr()
    pbdr = p_bdr.find(qn("w:pBdr"))
    if pbdr is None:
        pbdr = OxmlElement("w:pBdr")
        p_bdr.append(pbdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "3")
    bottom.set(qn("w:color"), BORDER)
    pbdr.append(bottom)


def add_title_page(doc):
    p = doc.add_paragraph()
    set_para_format(p, before=0, after=12, line=1.0)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    add_run_text(p, "CROUT HOLDINGS (PTY) LTD", size=12, color=GOLD, bold=True)

    title = doc.add_paragraph()
    set_para_format(title, before=28, after=8, line=1.0)
    add_run_text(title, "INDEPENDENT CONTRACTOR /", size=24, color=ACCENT_DARK, bold=True)
    add_run_text(title, "\nSUBCONTRACTOR AGREEMENT", size=24, color=ACCENT_DARK, bold=True)

    subtitle = doc.add_paragraph()
    set_para_format(subtitle, before=0, after=18, line=1.15)
    add_run_text(subtitle, "Crout Automations division | South African reusable commercial template", size=12.5, color=MUTED)

    intro_box = doc.add_table(rows=3, cols=1)
    intro_box.alignment = WD_TABLE_ALIGNMENT.LEFT
    intro_box.autofit = False
    intro_box.columns[0].width = Inches(6.5)
    items = [
        "Prepared for use with independent contractors and subcontractors engaged to deliver technology, automation, software, AI, integration, consulting and related digital services.",
        "This template is drafted to support an independent contractor relationship under South African law, including POPIA and electronic contracting provisions where relevant.",
        "Before production use, replace all placeholder fields and have the final version reviewed for the specific contractor, service profile and risk allocation.",
    ]
    for idx, text in enumerate(items):
        cell = intro_box.cell(idx, 0)
        cell.width = Inches(6.5)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_cell_shading(cell, "F6F8FA" if idx != 1 else "EEF3F7")
        set_cell_border(
            cell,
            top={"val": "single", "sz": 8, "color": BORDER},
            bottom={"val": "single", "sz": 8, "color": BORDER},
            left={"val": "single", "sz": 8, "color": BORDER},
            right={"val": "single", "sz": 8, "color": BORDER},
        )
        p = cell.paragraphs[0]
        set_para_format(p, before=0, after=0, line=1.15)
        add_run_text(p, text, size=10.5, color=TEXT)

    doc.add_paragraph("")

    meta = doc.add_table(rows=5, cols=2)
    meta.alignment = WD_TABLE_ALIGNMENT.LEFT
    meta.autofit = False
    widths = [Inches(1.85), Inches(4.65)]
    for row in meta.rows:
        row.cells[0].width = widths[0]
        row.cells[1].width = widths[1]
    details = [
        ("Principal", "Crout Holdings (Pty) Ltd"),
        ("Division", "Crout Automations"),
        ("Template version", "v1.0 | {{TemplateVersionDate}}"),
        ("Proposed contractor", "{{ContractorFullName}}"),
        ("Effective date", "{{StartDate}}"),
    ]
    for idx, (label, value) in enumerate(details):
        for col, text in enumerate((label, value)):
            cell = meta.cell(idx, col)
            set_cell_border(
                cell,
                top={"val": "single", "sz": "6", "color": BORDER},
                bottom={"val": "single", "sz": "6", "color": BORDER},
                left={"val": "single", "sz": "6", "color": BORDER},
                right={"val": "single", "sz": "6", "color": BORDER},
            )
            if col == 0:
                set_cell_shading(cell, "F3F5F7")
            p = cell.paragraphs[0]
            set_para_format(p, before=0, after=0, line=1.0)
            add_run_text(p, text, size=10.5, color=TEXT, bold=(col == 0))

    doc.add_page_break()


def add_toc_page(doc):
    label = doc.add_paragraph(style="Small Caps Label")
    add_run_text(label, "Document contents", size=9.5, color=MUTED)

    h = doc.add_paragraph(style="Heading 1")
    h.text = "Table of Contents"
    force_font(h.runs[0], name="Aptos", size=15, color=ACCENT, bold=True)

    entries = [
        ("Parties and Recitals", "3"),
        ("1. Definitions and Interpretation", "3"),
        ("2. Appointment and Nature of Engagement", "3"),
        ("3. Independent Contractor Status", "4"),
        ("4. Project Acceptance, Delivery and Support", "4"),
        ("5. Services and Performance Standards", "5"),
        ("6. Fees, Commission and Payment Mechanics", "5"),
        ("7. Invoicing, Records and Set-Off", "6"),
        ("8. Tax, Statutory and Business Compliance", "6"),
        ("9. Subcontracting and Contractor Teams", "7"),
        ("10. Confidential Information", "7"),
        ("11. Data Protection and POPIA", "8"),
        ("12. Information Security", "8"),
        ("13. Intellectual Property and Deliverables", "9"),
        ("14. Warranties, Indemnities and Liability", "9"),
        ("15. Non-Solicitation, Non-Circumvention and Conflict Management", "10"),
        ("16. Term and Termination", "10"),
        ("17. Consequences of Termination", "11"),
        ("18. Communications, E-Signatures and Notices", "11"),
        ("19. Dispute Resolution, Governing Law and General", "12"),
        ("20. Execution", "13"),
        ("Schedule A. Commercial Terms", "14"),
        ("Schedule B. Service Categories", "15"),
    ]
    for title_text, page in entries:
        p = doc.add_paragraph()
        set_para_format(p, before=0, after=2, line=1.05)
        tabs = p.paragraph_format.tab_stops
        tabs.add_tab_stop(Inches(6.15), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)
        add_run_text(p, title_text, size=10.7, color=TEXT)
        add_run_text(p, "\t")
        add_run_text(p, page, size=10.7, color=MUTED)
    doc.add_page_break()


def add_introductory_recitals(doc):
    heading = doc.add_paragraph(style="Heading 1")
    heading.text = "Parties and Recitals"
    force_font(heading.runs[0], size=15, color=ACCENT, bold=True)

    recitals = [
        ("A.", "Crout Holdings (Pty) Ltd is a South African company that offers technology-enabled services, including through its Crout Automations division."),
        ("B.", "The Company may from time to time source specialist digital services from independent contractors or subcontractors on a non-exclusive, project-by-project basis."),
        ("C.", "The Contractor has represented that the Contractor has the skill, capacity, tools and commercial independence necessary to provide the Services as an independent business undertaking."),
        ("D.", "The parties wish to record the terms on which the Contractor may accept and perform Services for the Company or the Company’s clients, while expressly avoiding the creation of an employment relationship."),
    ]
    for label, text in recitals:
        p = doc.add_paragraph(style="Clause List")
        add_run_text(p, f"{label} ", bold=True)
        add_run_text(p, text)


def add_clause_heading(doc, number, title):
    p = doc.add_paragraph(style="Heading 1")
    p.text = f"{number}. {title}"
    force_font(p.runs[0], size=15, color=ACCENT, bold=True)


def add_subheading(doc, text):
    p = doc.add_paragraph(style="Heading 2")
    p.text = text
    force_font(p.runs[0], size=12.5, color=ACCENT, bold=True)


def clause_paragraph(doc, text):
    p = doc.add_paragraph()
    set_para_format(p, before=0, after=6, line=1.12)
    add_run_text(p, text)


def clause_list(doc, items, numbered=True):
    style = "Clause List" if numbered else "List Bullet"
    for item in items:
        p = doc.add_paragraph(style=style)
        if style == "List Bullet":
            set_para_format(p, before=0, after=4, line=1.12)
            force_font(p.runs[0], size=11, color=TEXT) if p.runs else None
            add_run_text(p, item)
        else:
            add_run_text(p, item)


def add_main_clauses(doc):
    clauses = [
        (
            "Definitions and Interpretation",
            [
                "In this Agreement, unless the context indicates otherwise: \"Accepted Project\" means a discrete project, task order, work package, support assignment or client engagement that the Contractor has accepted in writing or through an agreed electronic workflow; \"Agreement\" means this Independent Contractor / Subcontractor Agreement, including its schedules and any written project acceptance; \"Client\" means any current or prospective customer of the Company in relation to whom Services may be performed; \"Company\" means Crout Holdings (Pty) Ltd; \"Contractor\" means {{ContractorFullName}} or the legal entity identified in Schedule A; \"Deliverables\" means all work product, materials, software, automation logic, documentation and outputs produced in connection with the Services; and \"Services\" means the service categories listed in Schedule B together with any related digital or technical services agreed by the parties.",
                "References to legislation include amendments, replacements and subordinate measures applicable in the Republic of South Africa from time to time. References to writing include electronic communications and data messages recognised under applicable law.",
            ],
            [],
        ),
        (
            "Appointment and Nature of Engagement",
            [
                "The Company appoints the Contractor, on a non-exclusive and as-needed basis, to perform Services for Accepted Projects, and the Contractor may accept those Services subject to this Agreement.",
                "Nothing in this Agreement obliges the Company to offer any minimum volume of work, retainer, exclusivity, revenue threshold or duration of engagement, and nothing obliges the Contractor to accept any offered project before the Contractor has expressly accepted it.",
            ],
            [
                "Each project opportunity may be offered by email, messaging platform, project board, ticketing tool, statement of work, quotation acceptance or other written instruction agreed by the parties.",
                "The Contractor becomes responsible for an Accepted Project only once the Contractor has confirmed acceptance in writing or through an agreed electronic system.",
                "The Contractor may carry on business for other clients, subject always to the confidentiality, conflict, data protection and non-solicitation restrictions in this Agreement.",
            ],
        ),
        (
            "Independent Contractor Status",
            [
                "The parties record and intend that the Contractor acts at all times as an independent contractor conducting an independent trade, occupation or profession, and not as an employee, worker, partner, agent in the labour law sense, or representative with authority to bind the Company unless separately authorised in writing.",
                "The Contractor retains control over the manner, methods, sequence, place and resources used to perform the Services, subject only to agreed specifications, delivery dates, security standards, quality requirements, lawful instructions relating to the Client environment, and the Company’s right to verify contractual compliance.",
            ],
            [
                "The Contractor is not entitled to paid leave, sick leave, annual leave, family responsibility leave, public holiday pay, overtime pay, severance pay, notice pay applicable to employees, bonuses, pension, provident fund participation, medical aid, UIF contributions from the Company or any other employee benefit.",
                "No provision of this Agreement creates any expectation of continued work, permanent engagement, fixed hours, supervision comparable to employment, or integration into the Company’s organisational structure as an employee.",
                "The Contractor is responsible for maintaining the Contractor’s own tools, devices, workspace, connectivity, utilities, licences, staffing and professional resources required to perform the Services, unless otherwise agreed in Schedule A for a specific Accepted Project.",
            ],
        ),
        (
            "Project Acceptance, Delivery and Support",
            [
                "The Contractor must review each proposed Accepted Project for scope, dependencies, delivery assumptions and commercial viability before acceptance.",
                "Unless otherwise agreed in writing for a specific Accepted Project, the Contractor must complete development work within a maximum of five calendar days after acceptance and complete testing, remediation and deployment readiness activities within a maximum of two additional calendar days, resulting in a total maximum delivery period of seven calendar days.",
            ],
            [
                "Time periods run from the date and time the Contractor accepts the project, provided the Company or Client has supplied the information and system access reasonably necessary to begin work.",
                "If the Contractor identifies a dependency, blocker or scope change that may affect delivery, the Contractor must notify the Company immediately and propose a revised plan in writing.",
                "After acceptance, the Contractor remains responsible for the completion of the Accepted Project, reasonable bug fixing, remediation and ongoing support obligations specifically attached to that Accepted Project, subject to the commercial terms in Schedule A or any later written variation.",
                "Repeated late delivery, avoidable defects, failure to communicate material risks, or failure to support an Accepted Project may be taken into account when the Company allocates future work.",
            ],
        ),
        (
            "Services and Performance Standards",
            [
                "The Services may include software development, automation development, artificial intelligence development, integrations, technical consulting, solution architecture, workflow design, quality assurance, support services and any other agreed digital services.",
                "The Contractor must perform the Services with the degree of skill, care, diligence, speed and professionalism reasonably expected from a suitably qualified independent specialist providing comparable services in South Africa.",
            ],
            [
                "The Contractor must comply with documented specifications, accepted scope, security requirements, development standards and applicable law relevant to the Services.",
                "The Contractor must keep adequate working records for each Accepted Project, including deliverable status, issues requiring escalation, credentials issued, third-party dependencies and material instructions received.",
                "No expansion of service categories under Schedule B requires a formal amendment if the added services are related digital, automation, technical or consulting services agreed in writing between the parties.",
            ],
        ),
        (
            "Fees, Commission and Payment Mechanics",
            [
                "As consideration for the Services, the Company must pay the Contractor the commission or other commercial amounts stated in Schedule A. Unless otherwise stated in Schedule A, the initial commission rate for Accepted Projects is {{CommissionPercentage}} of the relevant collected service revenue.",
                "Subject to this clause and Schedule A, the Contractor is entitled to recurring monthly commission for a completed service while the relevant Client remains active, the relevant service remains live and revenue-generating for the Company, and the Contractor remains entitled to payment under this Agreement.",
            ],
            [
                "Payment frequency is {{PaymentFrequency}} and payment must be made to the bank account stated in Schedule A against a valid invoice or remittance process reasonably required by the Company.",
                "Commission accrues only on revenue actually received by the Company from the Client in respect of the relevant service and does not accrue on cancelled work, refunded amounts, uncollected debt, write-offs, chargebacks, goodwill credits, taxes collected for remittance or pass-through third-party charges unless Schedule A expressly provides otherwise.",
                "Recurring commission ceases, without any further payment obligation, if the Client cancels or terminates the relevant service; the Client stops paying or becomes materially in arrears; the service is migrated away from the Company or replaced with another solution for which the Contractor is not contractually designated; the Company terminates the Contractor’s entitlement for breach under this Agreement; the service is lawfully discontinued; or the parties agree in writing to a different commercial arrangement.",
                "Unless Schedule A says otherwise, commission on a recurring service also ceases if the Contractor materially fails to provide agreed handover, maintenance, remediation or cooperation reasonably necessary to sustain that service after accepting responsibility for it.",
                "The Contractor bears all ordinary business expenses, including travel, devices, telecommunications, tax compliance, professional subscriptions and insurance, unless the Company expressly approves a reimbursable expense in writing before it is incurred.",
            ],
        ),
        (
            "Invoicing, Records and Set-Off",
            [
                "The Contractor must provide accurate invoices, statements or supporting information reasonably required by the Company to verify completed Services, commission calculations and tax treatment.",
                "The Company may withhold disputed amounts in good faith pending verification, provided it pays any undisputed portion when due.",
            ],
            [
                "If the Company is legally required to withhold or deduct tax or another statutory charge, the Company may do so and will account for the withheld amount as required by law.",
                "The Company may set off any liquidated amount owed by the Contractor to the Company against any amount otherwise due to the Contractor, including overpayments, chargebacks attributable to the Contractor’s breach, unauthorised Client refunds caused by the Contractor, or losses recoverable under an indemnity in this Agreement.",
            ],
        ),
        (
            "Tax, Statutory and Business Compliance",
            [
                "The Contractor is solely responsible for the Contractor’s own income tax, provisional tax, VAT registration and compliance where applicable, UIF obligations, PAYE obligations where legally applicable to the Contractor’s own personnel, accounting records, statutory registrations, exchange control compliance if relevant, and all filings, levies or regulatory obligations arising from the Contractor’s business.",
                "The Company does not assume responsibility for the Contractor’s payroll, tax administration, social security, employee deductions, workers compensation cover, public liability insurance, cyber insurance or other business insurance, unless expressly agreed in writing.",
            ],
            [
                "The Contractor must immediately notify the Company if the Contractor’s tax status, VAT status, banking details or legal identity changes.",
                "The Contractor warrants that the Contractor will comply with all laws applicable to the Contractor’s business operations and the performance of the Services, including but not limited to labour, immigration, intellectual property, privacy, anti-corruption, export control and consumer-protection related laws to the extent relevant.",
            ],
        ),
        (
            "Subcontracting and Contractor Teams",
            [
                "The Contractor may propose the use of the Contractor’s own personnel or downstream subcontractors to assist in performing the Services, but no such person may perform work, access repositories, interact with a Client, or access Company or Client systems unless first approved in writing by the Company.",
                "Each proposed team member or subcontractor must complete any interview, due diligence, skill assessment, confidentiality process, data protection onboarding and security review reasonably required by the Company before approval is granted.",
            ],
            [
                "The Contractor remains fully responsible for the acts, omissions, quality, security compliance, timing, cost and legal compliance of every approved team member or subcontractor used by the Contractor.",
                "The Contractor must ensure that each approved team member or subcontractor is bound by written obligations no less protective than those contained in this Agreement, including obligations relating to confidentiality, POPIA, information security, intellectual property assignment and non-solicitation.",
                "Approval of a subcontractor may be withheld or withdrawn by the Company in its reasonable discretion if risk, performance or Client requirements justify that decision.",
            ],
        ),
        (
            "Confidential Information",
            [
                "The Contractor must keep strictly confidential all Confidential Information obtained from or relating to the Company, its affiliates, its Clients or their respective business affairs, whether disclosed before or after the Effective Date and whether in oral, written, visual, electronic or other form.",
                "Confidential Information includes client data, personal information, pricing, proposals, financial information, repositories, source code, object code, scripts, prompts, models, automation logic, APIs, credentials, infrastructure details, business processes, workflows, standard operating procedures, security measures, roadmaps, know-how, documentation, trade secrets and any information marked or reasonably understood to be confidential.",
            ],
            [
                "The Contractor may use Confidential Information only to perform Accepted Projects and may not disclose it to any third party except to an approved subcontractor with a strict need to know and an equivalent written confidentiality obligation.",
                "The Contractor must protect Confidential Information with at least reasonable care and, in any event, not less than the care the Contractor uses to protect its own most sensitive information.",
                "The obligations in this clause continue for the duration of the Agreement and indefinitely thereafter in relation to trade secrets, proprietary source code, credentials and personal information, and for not less than five years after termination in relation to other Confidential Information unless a longer period is required by law or another written obligation.",
            ],
        ),
        (
            "Data Protection and POPIA",
            [
                "To the extent the Contractor processes personal information for or on behalf of the Company or a Client, the Contractor must do so only on documented instructions, only for the purpose of performing the relevant Accepted Project, and in compliance with the Protection of Personal Information Act, 2013 and any binding code, guidance or lawful operator instructions applicable to that processing.",
                "The Contractor must implement appropriate technical and organisational measures to secure the integrity and confidentiality of personal information against loss, damage, unauthorised destruction, unlawful access or unlawful processing, taking into account generally accepted information security practices and the sensitivity of the data concerned.",
            ],
            [
                "The Contractor may not sell, mine, train models on, commercialise, combine or repurpose personal information or Client data except to the extent expressly authorised in writing by the Company and lawful under POPIA.",
                "The Contractor must notify the Company immediately, and in any event without undue delay, after becoming aware of any actual or suspected personal information breach, unauthorised disclosure, credential compromise, unlawful access event or other security incident affecting Company or Client data.",
                "The Contractor must cooperate fully with the Company in investigating, containing, remediating and, where required, notifying affected parties or regulators in connection with any data or security incident.",
                "On request or on termination, the Contractor must return or securely delete personal information and other regulated data in the Contractor’s possession or control, except to the extent retention is required by law, in which case the retained information must remain protected under this Agreement.",
            ],
        ),
        (
            "Information Security",
            [
                "The Contractor must maintain appropriate security controls for all systems, devices, accounts and environments used to perform the Services, including strong passwords, multi-factor authentication where available, device encryption where reasonably appropriate, secure credential storage, timely patching, controlled access management and secure backup practices.",
                "The Contractor must use only approved repositories, storage locations, collaboration channels and deployment pathways for Company or Client work, unless the Company gives prior written consent to an alternative arrangement.",
            ],
            [
                "The Contractor must not share credentials, disable security tooling, copy data to unmanaged devices, or permit any unauthorised person to access Company or Client systems.",
                "The Contractor must keep development environments, logs, exports and local copies of data limited to what is reasonably necessary and must promptly remove obsolete access and stale data copies.",
                "Any loss of a device, credential compromise, malware event, suspected account takeover or other security incident relevant to the Services must be reported to the Company immediately.",
            ],
        ),
        (
            "Intellectual Property and Deliverables",
            [
                "All Deliverables, inventions, discoveries, enhancements, configurations, source code, object code, documentation, databases, APIs, integrations, prompts, training materials, scripts, workflows, process maps, SOPs, internal tools, designs, reports and other materials created, developed, authored, adapted or reduced to practice by or for the Contractor in connection with the Services vest in and are assigned to the Company, to the fullest extent permitted by law, upon creation and, where necessary to perfect title, irrevocably on the later of creation and payment for the relevant Deliverable.",
                "To the extent any right, including copyright, design right, database right, confidential information right, neighbouring right or similar intellectual property right does not vest automatically, the Contractor hereby irrevocably cedes, assigns and transfers that right to the Company with full title guarantee, including the right to claim and recover for past infringements insofar as lawful.",
            ],
            [
                "The Contractor waives, to the fullest extent permitted by law, any moral rights or similar rights that may conflict with the Company’s full use, adaptation, publication, modification or commercialisation of the Deliverables, and where waiver is not legally effective the Contractor irrevocably consents to those acts.",
                "The Contractor must execute all documents and do all acts reasonably required by the Company to evidence, perfect, register or enforce the Company’s ownership of the Deliverables or related intellectual property rights.",
                "The Contractor may not reuse, license, resell or disclose Deliverables created for the Company or a Client, except for residual know-how and general skills that do not reveal Confidential Information, personal information or proprietary code and do not infringe the Company’s ownership rights.",
                "If the Contractor incorporates any third-party materials, open-source components or pre-existing tools into a Deliverable, the Contractor must disclose that fact in advance, ensure the relevant licence permits the intended use, and not include any material that imposes unexpected disclosure, copyleft, access, royalty or downstream restriction obligations on the Company or its Client without prior written approval.",
            ],
        ),
        (
            "Warranties, Indemnities and Liability",
            [
                "The Contractor warrants that the Services and Deliverables will materially conform to agreed specifications, be performed with reasonable skill and care, and not knowingly infringe the intellectual property rights or privacy rights of any third party.",
                "The Contractor indemnifies the Company and its Clients against direct loss, liability, claim, cost and damage arising from the Contractor’s breach of confidentiality, data protection obligations, unlawful conduct, infringement of intellectual property rights, unauthorised subcontracting, fraud, wilful misconduct or gross negligence.",
            ],
            [
                "Neither party excludes liability for fraud, wilful misconduct, gross negligence, death or personal injury caused by negligence, breach of confidentiality, breach of data protection obligations or any liability that cannot lawfully be excluded or limited.",
                "Subject to the previous item, the Company’s aggregate liability to the Contractor arising out of this Agreement is limited to the total fees or commission actually paid or payable to the Contractor for the Accepted Project giving rise to the claim during the six months preceding the event giving rise to the claim.",
                "The Company is not liable for indirect, consequential, special or pure economic loss, loss of profit, loss of opportunity or reputational harm suffered by the Contractor, except to the extent such exclusion is unlawful.",
            ],
        ),
        (
            "Non-Solicitation, Non-Circumvention and Conflict Management",
            [
                "During the term of this Agreement and for {{NonSolicitPeriod}} after termination, the Contractor must not, without the Company’s prior written consent, directly or indirectly solicit, entice away, divert, bypass or contract directly with any Client or active prospect introduced by or materially serviced through the Company in relation to services that compete with or substitute for the Services performed under this Agreement.",
                "During the same period, the Contractor must not knowingly solicit for employment or engagement any employee, core contractor or approved subcontractor of the Company with whom the Contractor had material dealings through this Agreement.",
            ],
            [
                "This clause does not prohibit responses to a genuine public recruitment campaign or general advertising not specifically targeted at protected persons, provided there is no deliberate circumvention of this Agreement.",
                "The Contractor must disclose any actual or reasonably apparent conflict of interest that could affect impartial performance, data security, independence or the Company’s client relationships.",
                "The parties acknowledge that the restrictions in this clause are intended to protect legitimate proprietary interests and must be interpreted to be enforceable to the maximum extent permitted under South African law.",
            ],
        ),
        (
            "Term and Termination",
            [
                "This Agreement starts on {{StartDate}} and continues until terminated in accordance with this clause.",
                "Either party may terminate this Agreement for convenience by giving the other party {{NoticePeriod}} written notice, provided that termination for convenience does not relieve the Contractor from completing or properly handing over any Accepted Project already in progress unless the Company directs otherwise.",
            ],
            [
                "The Company may terminate this Agreement immediately by written notice if the Contractor commits a material breach, including a breach of confidentiality, POPIA, security obligations, intellectual property obligations, non-solicitation obligations, fraud, dishonesty, gross misconduct, unlawful conduct, unauthorised subcontracting, repeated failure to deliver Accepted Projects or conduct that materially endangers a Client relationship.",
                "The Contractor may terminate immediately if the Company commits a material breach and fails to remedy that breach within ten business days after written demand, provided the breach is capable of remedy.",
                "On termination, the Contractor is entitled only to amounts properly due for Services validly performed up to the termination date, less lawful deductions, set-off and amounts withheld pending verification or remediation.",
                "If termination occurs for the Contractor’s breach, the Company may withhold commission on incomplete or defective work until the loss, rework cost, refund exposure or handover impact has been quantified.",
            ],
        ),
        (
            "Consequences of Termination",
            [
                "Immediately on expiry or termination, the Contractor must stop representing any connection with the Company except as necessary to complete an agreed handover, cease using Company and Client systems, return or securely destroy Company property, and deliver all work in progress, credentials, access notes, repositories, project records and other materials relating to the Services.",
                "At the Company’s request, the Contractor must certify in writing that all Confidential Information, personal information, credentials and Company property in the Contractor’s possession or control have been returned, revoked or securely deleted, subject to any lawful retention requirement.",
            ],
            [
                "Termination does not affect accrued rights, accepted payment obligations already due, or clauses that by their nature are intended to survive termination, including confidentiality, POPIA, security, intellectual property, indemnities, non-solicitation, dispute resolution and governing law.",
                "If an Accepted Project is transitioned away from the Contractor, the Contractor must provide reasonable handover assistance at the rates or commission adjustment stated in Schedule A, or if no rate is stated, at a reasonable rate agreed in writing.",
            ],
        ),
        (
            "Communications, E-Signatures and Notices",
            [
                "The parties agree that project instructions, approvals, notices, variations, acceptance messages, invoices and other operational communications may be transmitted by email or other agreed electronic channel, and such communications may constitute valid data messages and records for the purposes of applicable law.",
                "This Agreement and any project acceptance may be signed electronically, including by advanced electronic signature where required by law, counterpart signature, scanned signature page or platform-based signature process accepted by the parties, and each signed counterpart forms part of one agreement.",
            ],
            [
                "Formal notices under this Agreement must be sent to the addresses and email contacts stated in Schedule A, unless updated by written notice.",
                "A notice is deemed received on the date of delivery if delivered by hand on a business day, or on transmission if sent by email before 17:00 on a business day to the nominated email address and no delivery failure message is received, failing which on the next business day.",
            ],
        ),
        (
            "Dispute Resolution, Governing Law and General",
            [
                "This Agreement is governed by the laws of the Republic of South Africa.",
                "If a dispute arises out of or in connection with this Agreement, a party must first give written notice of the dispute and the parties must seek to resolve it through good-faith senior-level negotiations within ten business days.",
            ],
            [
                "If the dispute is not resolved through negotiation, the parties must attempt mediation in Johannesburg, South Africa, using a mediator agreed between them or, failing agreement, nominated by the Arbitration Foundation of Southern Africa or its successor.",
                "If mediation does not resolve the dispute within fifteen business days after appointment of the mediator, either party may refer the dispute to arbitration in Johannesburg under AFSA’s commercial rules, provided either party may seek urgent interim relief from a court of competent jurisdiction.",
                "This Agreement constitutes the entire agreement between the parties regarding its subject matter and supersedes prior proposals or discussions relating to the same subject matter.",
                "No amendment or waiver is effective unless in writing and signed or otherwise validly accepted by the parties through an agreed legally recognised method.",
                "If any provision is found unenforceable, it must be severed or read down only to the minimum extent necessary without affecting the enforceability of the remaining provisions.",
            ],
        ),
    ]

    for idx, (title, paragraphs, list_items) in enumerate(clauses, start=1):
        add_clause_heading(doc, idx, title)
        for paragraph in paragraphs:
            clause_paragraph(doc, paragraph)
        if list_items:
            clause_list(doc, list_items, numbered=True)


def add_signature_page(doc):
    doc.add_page_break()
    add_clause_heading(doc, 20, "Execution")
    clause_paragraph(
        doc,
        "The parties confirm that they have read, understood and agree to be bound by this Agreement, including Schedules A and B, from the Effective Date.",
    )

    table = doc.add_table(rows=7, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    widths = [Inches(3.15), Inches(3.15)]
    for row in table.rows:
        row.cells[0].width = widths[0]
        row.cells[1].width = widths[1]

    headings = [
        "For Crout Holdings (Pty) Ltd",
        "For the Contractor",
    ]
    for col, text in enumerate(headings):
        cell = table.cell(0, col)
        set_cell_shading(cell, "F3F5F7")
        set_cell_border(
            cell,
            top={"val": "single", "sz": "10", "color": BORDER},
            bottom={"val": "single", "sz": "10", "color": BORDER},
            left={"val": "single", "sz": "10", "color": BORDER},
            right={"val": "single", "sz": "10", "color": BORDER},
        )
        p = cell.paragraphs[0]
        set_para_format(p, before=0, after=0, line=1.0)
        add_run_text(p, text, size=10.5, color=ACCENT_DARK, bold=True)

    rows = [
        ("Name: {{CompanySignatoryName}}", "Name: {{ContractorFullName}}"),
        ("Capacity: {{CompanySignatoryCapacity}}", "ID / Reg No.: {{IDNumber}}"),
        ("Signature: ____________________", "Signature: ____________________"),
        ("Date: {{SignedDate}}", "Date: {{SignedDate}}"),
        ("Place: {{SignaturePlace}}", "Place: {{SignaturePlace}}"),
        ("Email: {{CompanySignatoryEmail}}", "Email: {{Email}}"),
    ]
    for row_idx, (left, right) in enumerate(rows, start=1):
        for col, text in enumerate((left, right)):
            cell = table.cell(row_idx, col)
            set_cell_border(
                cell,
                top={"val": "single", "sz": "6", "color": BORDER},
                bottom={"val": "single", "sz": "6", "color": BORDER},
                left={"val": "single", "sz": "6", "color": BORDER},
                right={"val": "single", "sz": "6", "color": BORDER},
            )
            p = cell.paragraphs[0]
            set_para_format(p, before=0, after=0, line=1.0)
            add_run_text(p, text, size=10.5, color=TEXT)


def add_schedule_a(doc):
    doc.add_page_break()
    add_clause_heading(doc, "Schedule A", "Commercial Terms")
    clause_paragraph(
        doc,
        "This Schedule A forms part of the Agreement and may be updated for each contractor or service profile without amending the operative clauses, provided the update is agreed in writing.",
    )

    table = doc.add_table(rows=11, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    widths = [Inches(2.1), Inches(4.4)]
    for row in table.rows:
        row.cells[0].width = widths[0]
        row.cells[1].width = widths[1]
    entries = [
        ("Contractor legal name", "{{ContractorFullName}}"),
        ("ID / registration number", "{{IDNumber}}"),
        ("Email address", "{{Email}}"),
        ("Physical address", "{{PhysicalAddress}}"),
        ("Effective date", "{{StartDate}}"),
        ("Primary service type", "{{ServiceType}}"),
        ("Commission percentage", "{{CommissionPercentage}}"),
        ("Payment frequency", "{{PaymentFrequency}}"),
        ("Bank name / branch", "{{BankName}} / {{BankBranch}}"),
        ("Account details", "{{BankAccountNumber}} / {{BankAccountType}}"),
        ("Notice period / special terms", "{{NoticePeriod}} / {{SpecialCommercialTerms}}"),
    ]
    for row_idx, (label, value) in enumerate(entries):
        label_cell = table.cell(row_idx, 0)
        value_cell = table.cell(row_idx, 1)
        set_cell_shading(label_cell, "F7F8FA")
        for cell in (label_cell, value_cell):
            set_cell_border(
                cell,
                top={"val": "single", "sz": "6", "color": BORDER},
                bottom={"val": "single", "sz": "6", "color": BORDER},
                left={"val": "single", "sz": "6", "color": BORDER},
                right={"val": "single", "sz": "6", "color": BORDER},
            )
        p1 = label_cell.paragraphs[0]
        set_para_format(p1, before=0, after=0, line=1.0)
        add_run_text(p1, label, size=10.5, color=TEXT, bold=True)
        p2 = value_cell.paragraphs[0]
        set_para_format(p2, before=0, after=0, line=1.0)
        add_run_text(p2, value, size=10.5, color=TEXT)

    add_subheading(doc, "Recurring Commission Guidance")
    clause_list(
        doc,
        [
            "Recurring commission remains payable only while the relevant client account is active and collected, the relevant service remains in force, and the Contractor remains entitled under the Agreement.",
            "Examples of cessation events include client cancellation, non-payment, insolvency, replacement of the service, migration to another provider, service discontinuance, breach-based termination, or a revised commercial structure accepted in writing.",
            "If a project has separate setup, support or maintenance economics, describe those mechanics here: {{RecurringCommissionNotes}}.",
        ],
        numbered=True,
    )


def add_schedule_b(doc):
    doc.add_page_break()
    add_clause_heading(doc, "Schedule B", "Service Categories")
    clause_paragraph(
        doc,
        "The following categories are examples of Services that may be allocated to the Contractor. The list is illustrative and may be expanded by written agreement without amending the core Agreement.",
    )

    table = doc.add_table(rows=8, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    widths = [Inches(2.2), Inches(4.3)]
    for row in table.rows:
        row.cells[0].width = widths[0]
        row.cells[1].width = widths[1]

    services = [
        ("Software Development", "Web applications, APIs, backend services, product features, scripts and custom tooling."),
        ("Automation Development", "Workflow automation, CRM processes, no-code or low-code systems, task orchestration and internal automations."),
        ("AI Development", "Prompt engineering, AI workflows, agent systems, integrations with foundation models and applied AI implementation."),
        ("Integrations", "Third-party APIs, data pipelines, middleware, webhooks and systems integration work."),
        ("Technical Consulting", "Architecture, scoping, technical advisory, troubleshooting, audits and implementation guidance."),
        ("Support and Maintenance", "Bug fixes, service continuity tasks, optimisation and agreed support services for Accepted Projects."),
        ("Documentation and SOPs", "Technical documentation, client handover packs, workflows, process notes and system operating guides."),
        ("Other Agreed Services", "{{AdditionalServiceCategories}}"),
    ]

    for idx, (service, description) in enumerate(services):
        for col, text in enumerate((service, description)):
            cell = table.cell(idx, col)
            if idx == 0:
                set_repeat_table_header(table.rows[0])
            if col == 0:
                set_cell_shading(cell, "F7F8FA")
            set_cell_border(
                cell,
                top={"val": "single", "sz": "6", "color": BORDER},
                bottom={"val": "single", "sz": "6", "color": BORDER},
                left={"val": "single", "sz": "6", "color": BORDER},
                right={"val": "single", "sz": "6", "color": BORDER},
            )
            p = cell.paragraphs[0]
            set_para_format(p, before=0, after=0, line=1.0)
            add_run_text(p, text, size=10.3, color=TEXT, bold=(col == 0))


def add_final_note(doc):
    doc.add_page_break()
    add_subheading(doc, "Template Completion Checklist")
    clause_list(
        doc,
        [
            "Replace every placeholder field enclosed in double braces before issue.",
            "Confirm the Contractor’s legal identity, banking details, VAT status, notice period and commission model in Schedule A.",
            "Tailor any project-specific handover, maintenance or support obligations if the engagement differs from the default seven-day delivery framework.",
            "Have the execution-ready version reviewed against the actual contractor structure, client risk profile and any sector-specific regulatory constraints.",
        ],
        numbered=True,
    )


def build_document():
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document()
    set_doc_language(doc)
    configure_styles(doc)

    for section in doc.sections:
        set_section_layout(section)
        add_header(section)
        add_footer(section)

    add_title_page(doc)
    add_toc_page(doc)
    add_introductory_recitals(doc)
    add_main_clauses(doc)
    add_signature_page(doc)
    add_schedule_a(doc)
    add_schedule_b(doc)
    add_final_note(doc)

    doc.save(OUTPUT_DOCX)
    return OUTPUT_DOCX


if __name__ == "__main__":
    path = build_document()
    print(path)
